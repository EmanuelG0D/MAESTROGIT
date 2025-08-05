#EMANUEL HERNANDO HIGUERA VANEGAS
#05/08/2025
#WPOSS


# --- Multi-Convertidor Universal de Archivos ---
import os
import tkinter as tk
from tkinter import filedialog, messagebox
from docx import Document
from fpdf import FPDF
import pdfplumber
import markdown

# --- Funciones de conversión ---
def txt_to_docx(input_path, output_path):
    with open(input_path, 'r', encoding='utf-8') as f:
        lines = [line.rstrip() for line in f]
    doc = Document()
    for line in lines:
        doc.add_paragraph(line)
    doc.save(output_path)

def txt_to_pdf(input_path, output_path):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font('Arial', '', 12)
    with open(input_path, 'r', encoding='utf-8') as f:
        for line in f:
            pdf.cell(0, 10, line.rstrip(), ln=1)
    pdf.output(output_path)

def pdf_to_txt(input_path, output_path):
    with pdfplumber.open(input_path) as pdf:
        text = ''
        for page in pdf.pages:
            text += page.extract_text() + '\n'
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(text)

def txt_to_md(input_path, output_path):
    with open(input_path, 'r', encoding='utf-8') as f:
        content = f.read()
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(content)

def md_to_pdf(input_path, output_path):
    with open(input_path, 'r', encoding='utf-8') as f:
        html = markdown.markdown(f.read())
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font('Arial', '', 12)
    # Simple render: remove HTML tags for plain text
    import re
    text = re.sub('<[^<]+?>', '', html)
    for line in text.split('\n'):
        pdf.cell(0, 10, line, ln=1)
    pdf.output(output_path)

def docx_to_txt(input_path, output_path):
    doc = Document(input_path)
    text = '\n'.join([p.text for p in doc.paragraphs])
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(text)


# --- Mapeo de conversiones posibles ---
CONVERSION_FUNCTIONS = {
    ('txt', 'docx'): txt_to_docx,
    ('txt', 'pdf'): txt_to_pdf,
    ('txt', 'md'): txt_to_md,
    ('pdf', 'txt'): pdf_to_txt,
    ('md', 'pdf'): md_to_pdf,
    ('docx', 'txt'): docx_to_txt,
}

EXTENSIONS = {
    'TXT': '.txt',
    'DOCX': '.docx',
    'PDF': '.pdf',
    'MD': '.md',
}

# --- Interfaz gráfica ---

def main():
    root = tk.Tk()
    root.title('🔄 Multi-Convertidor Universal de Archivos')
    root.geometry('440x320')
    root.configure(bg='#f4f6fb')
    root.resizable(False, False)

    tk.Label(root, text='Multi-Convertidor Universal', font=('Segoe UI', 16, 'bold'), bg='#f4f6fb', fg='#2d3a4a').pack(pady=(18, 2))
    tk.Label(root, text='Convierte entre múltiples formatos de texto y documentos', font=('Segoe UI', 10), bg='#f4f6fb', fg='#4a5a6a').pack(pady=(0, 12))

    frame = tk.Frame(root, bg='#f4f6fb')
    frame.pack(pady=5)

    tk.Label(frame, text='Tipo de archivo origen:', font=('Segoe UI', 11), bg='#f4f6fb').grid(row=0, column=0, padx=8, pady=4, sticky='e')
    tk.Label(frame, text='Tipo de archivo destino:', font=('Segoe UI', 11), bg='#f4f6fb').grid(row=1, column=0, padx=8, pady=4, sticky='e')

    origen_var = tk.StringVar(value='TXT')
    destino_var = tk.StringVar(value='PDF')

    def actualizar_destinos(*args):
        origen = origen_var.get().lower()
        opciones = [k for k in EXTENSIONS.keys() if (origen, k.lower()) in CONVERSION_FUNCTIONS and k != origen_var.get()]
        menu = destino_menu['menu']
        menu.delete(0, 'end')
        for op in opciones:
            menu.add_command(label=op, command=tk._setit(destino_var, op))
        if opciones:
            destino_var.set(opciones[0])
        else:
            destino_var.set('')

    origen_menu = tk.OptionMenu(frame, origen_var, *EXTENSIONS.keys(), command=lambda _: actualizar_destinos())
    origen_menu.config(width=10, font=('Segoe UI', 11))
    origen_menu.grid(row=0, column=1, padx=8, pady=4)

    destino_menu = tk.OptionMenu(frame, destino_var, '')
    destino_menu.config(width=10, font=('Segoe UI', 11))
    destino_menu.grid(row=1, column=1, padx=8, pady=4)

    actualizar_destinos()

    archivo_var = tk.StringVar()

    def seleccionar_archivo():
        ext = EXTENSIONS[origen_var.get()]
        filetypes = [(f'Archivos {ext.upper()}', f'*{ext}'), ("Todos", "*.*")]
        path = filedialog.askopenfilename(title='Selecciona archivo de entrada', filetypes=filetypes)
        archivo_var.set(path)

    tk.Button(root, text='Seleccionar archivo de entrada', command=seleccionar_archivo, font=('Segoe UI', 11), bg='#e0e7ef', fg='#2d3a4a', relief='flat').pack(pady=(12, 2))
    tk.Entry(root, textvariable=archivo_var, width=48, font=('Segoe UI', 10), state='readonly', relief='groove').pack(pady=(0, 8))

    def convertir():
        origen = origen_var.get().lower()
        destino = destino_var.get().lower()
        input_path = archivo_var.get()
        if not input_path:
            messagebox.showerror('Error', 'Selecciona un archivo de entrada.')
            return
        if (origen, destino) not in CONVERSION_FUNCTIONS:
            messagebox.showerror('Error', 'Conversión no soportada.')
            return
        base = os.path.splitext(os.path.basename(input_path))[0]
        desktop = os.path.join(os.path.expanduser('~'), 'Desktop')
        output_path = os.path.join(desktop, base + EXTENSIONS[destino.upper()])
        try:
            CONVERSION_FUNCTIONS[(origen, destino)](input_path, output_path)
            messagebox.showinfo('Éxito', f'Archivo convertido guardado en:\n{output_path}')
            root.destroy()
        except Exception as e:
            messagebox.showerror('Error', str(e))

    tk.Button(root, text='Convertir', command=convertir, bg='#4CAF50', fg='white', font=('Segoe UI', 13, 'bold'), relief='flat', height=1, width=16).pack(pady=16)

    tk.Label(root, text='Desarrollado por tu asistente IA', font=('Segoe UI', 8), bg='#f4f6fb', fg='#7a8a9a').pack(side='bottom', pady=6)

    root.mainloop()

if __name__ == '__main__':
    main()
